import glob
import json
import os
import re
import subprocess
import sys
import time
import urllib.error
import urllib.request
from docx import Document
from gemini_utils import (
    RESPONSE_SELECTOR,
    find_input_box,
    find_send_button,
    select_gemini_model,
    start_clean_gemini_chat,
    wait_for_gemini_response,
)
from playwright.sync_api import sync_playwright
from utils import (
    get_config_value,
    kill_cdp_chrome,
    launch_browser_with_profile,
    rotate_profile_index,
    send_telegram_notification,
)

sys.stdout.reconfigure(encoding="utf-8")


def get_latest_run_folder(runs_path="youtube_runs"):
    if not os.path.exists(runs_path):
        return None
    folders = glob.glob(os.path.join(runs_path, "*/"))
    if not folders:
        return None
    return max(folders, key=os.path.getmtime)


def read_refine_prompt():
    try:
        with open("refine_prompt.txt", "r", encoding="utf-8") as f:
            return f.read().strip()
    except FileNotFoundError:
        print("Error: 'refine_prompt.txt' not found.")
        sys.exit(1)


def read_final_output(folder):
    path = os.path.join(folder, "final_output.txt")
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except FileNotFoundError:
        print(f"Error: 'final_output.txt' not found in {folder}")
        sys.exit(1)


def split_paragraphs(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if len(paragraphs) < 2:
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    return paragraphs


def load_checkpoint(folder):
    path = os.path.join(folder, "refine_checkpoint.json")
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("refined_paragraphs", [])
        except Exception:
            pass
    return []


def save_checkpoint(folder, refined_paragraphs):
    path = os.path.join(folder, "refine_checkpoint.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(
            {"refined_paragraphs": refined_paragraphs},
            f,
            ensure_ascii=False,
            indent=2,
        )


def delete_checkpoint(folder):
    path = os.path.join(folder, "refine_checkpoint.json")
    if os.path.exists(path):
        os.remove(path)

def apply_tashkeel_from_config(text):
    """Automatically replaces ambiguous Egyptian slang with correct Tashkeel."""
    config_path = "daheeh_config.json"
    if not os.path.exists(config_path):
        return text

    with open(config_path, "r", encoding="utf-8") as f:
        config = json.load(f)

    lexicon = (
        config.get("al_daheeh_master_pipeline_config", {})
        .get("dialect_profile", {})
        .get("tashkeel_lexicon", {})
    )

    # Automatically replace words like 'كده' with 'كِدَه'
    for plain_word, vocalized_word in lexicon.items():
        # Match standalone word boundaries
        pattern = rf"\b{re.escape(plain_word)}\b"
        text = re.sub(pattern, vocalized_word, text)

    return text

def clean_refined_paragraph(text):
    """Clean Gemini response and apply phonetic Tashkeel for TTS."""
    if not text:
        return ""

    # Strip XML tags & ledgers
    text = re.sub(
        r"<thinking>.*?</thinking>", "", text, flags=re.DOTALL | re.IGNORECASE
    )
    text = re.sub(
        r"<slang_ledger>.*?</slang_ledger>",
        "",
        text,
        flags=re.DOTALL | re.IGNORECASE,
    )
    text = re.sub(r"<.*?>", "", text)

    # Strip leaked English metadata blocks
    text = re.sub(
        r"(?:BLACKLIST|BRAINSTORM|UPDATE LEDGER|DRAFT & VERIFY).*?(?:Tashkeel added[^\.\n]*[\.\s]*|verified[\.\s]*|words[\.\s]*)",
        "",
        text,
        flags=re.DOTALL | re.IGNORECASE,
    )

    # Strip markdown and clean whitespace
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    cleaned_lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip() and not line.strip().startswith("BLACKLIST:")
    ]
    cleaned = " ".join(cleaned_lines)
    cleaned = re.sub(r"\([A-Za-z0-9\s\-_,\.\'&]+\)", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()

    # ◄── Apply Tashkeel from daheeh_config.json right here ──►
    return apply_tashkeel_from_config(cleaned)

def save_refined_script(folder, refined_paragraphs):
    text_path = os.path.join(folder, "refined_script.txt")
    docx_path = os.path.join(folder, "refined_script.docx")
    full_text = "\n\n".join(refined_paragraphs)

    with open(text_path, "w", encoding="utf-8") as f:
        f.write(full_text)

    doc = Document()
    doc.add_heading("Refined Script", level=1)
    for p in refined_paragraphs:
        doc.add_paragraph(p)
    doc.save(docx_path)

    print(f"Refined script saved: {text_path}")


def is_safety_blocked(text):
    if not text or len(text.strip()) < 15:
        return True
    lower_text = text.lower()
    refusal_keywords = [
        "cannot fulfill",
        "unable to assist",
        "safety guidelines",
        "against my policy",
        "something went wrong",
        "restricted content",
        "i am unable",
        "i apologize, but i cannot",
        "as an ai language model",
    ]
    for word in refusal_keywords:
        if word in lower_text:
            return True
    return False


def ensure_chrome_debug_session(browser_type, profile_index):
    """Verify or launch Chrome debugging session using the specified profile index from config."""
    url = "http://127.0.0.1:9222/json/version"

    # Check if port 9222 is open
    try:
        with urllib.request.urlopen(url, timeout=2) as response:
            if response.status == 200:
                print(f"Chrome debugging session is already active on port 9222.")
                return True
    except Exception:
        pass

    print(
        f"Launching Chrome debugging session (Browser: {browser_type}, Profile Index: {profile_index})..."
    )
    return launch_browser_with_profile(browser_type, profile_index)


def setup_refinement_session(page, model_name):
    """Turn 1: Send the refinement style guide and wait for acknowledgment."""
    print("[SETUP] Starting refinement session...")
    start_clean_gemini_chat(page)
    time.sleep(2)

    select_gemini_model(page, model_name)
    time.sleep(2)

    prompt_text = read_refine_prompt()

    print("[SETUP] Sending refinement style guide...")
    input_box = find_input_box(page)
    if not input_box:
        print("[FATAL] Could not find input box for setup turn.")
        return False

    input_box.click()
    time.sleep(0.5)
    input_box.fill(prompt_text)
    time.sleep(1)

    send_btn = find_send_button(page)
    if send_btn:
        send_btn.click()
    else:
        page.keyboard.press("Enter")

    print("[SETUP] Waiting for Gemini acknowledgment...")
    initial_count = page.locator(RESPONSE_SELECTOR).count()
    response = wait_for_gemini_response(
        page, initial_count, timeout_seconds=120
    )

    if response and (
        "understood" in response.lower()
        or "جاهز" in response.lower()
        or "مستعد" in response.lower()
    ):
        print("[SETUP] Gemini acknowledged refinement rules.")
        return True
    else:
        print(
            f"[SETUP] Gemini response: {response[:200] if response else '(empty)'}"
        )
        print(
            "[SETUP] Proceeding anyway — Gemini may have acknowledged implicitly."
        )
        return True


def refine_paragraph(page, paragraph_text, index, total):
    """Turn 2: Send a single paragraph for refinement."""
    persona = get_config_value("REFINE_PERSONA", "Al-Daheeh")
    max_slang = get_config_value("REFINE_MAX_SLANG_PER_SENTENCE", "2")
    
    message = (
        f"Refine paragraph {index} of {total}.\n"
        "SYSTEM OVERRIDE REMINDER:\n"
        f'1. Maintain "{persona}" persona (30% Fusha Data : 70% Cairene Amiya).\n'
        f"2. Follow the 1-3-1 sentence length cadence (Writing Music).\n"
        "3. Ground all facts in everyday Egyptian archetypes (Bureaucracy/Street Food/Installments).\n"
        "4. Apply targeted Tashkeel to ambiguous colloquial words.\n"
        "5. Check your previous <slang_ledger> and do not repeat words.\n\n"
        f"REFINE THIS PARAGRAPH:\n{paragraph_text}"
    )

    input_box = find_input_box(page)
    if not input_box:
        print(f"[ERROR] Could not find input box for paragraph {index}.")
        return None

    input_box.click()
    time.sleep(0.5)
    input_box.fill(message)
    time.sleep(1)

    send_btn = find_send_button(page)
    if send_btn:
        send_btn.click()
    else:
        page.keyboard.press("Enter")

    refine_timeout = int(get_config_value("REFINE_PARAGRAPH_TIMEOUT", "300"))
    initial_count = page.locator(RESPONSE_SELECTOR).count()
    response = wait_for_gemini_response(
        page, initial_count, timeout_seconds=refine_timeout
    )

    if is_safety_blocked(response):
        print(
            f"[WARNING] Safety block detected for paragraph {index}. Retrying with fresh chat..."
        )
        return None

    if response:
        # Apply strict cleaning & paragraph flattening
        return clean_refined_paragraph(response)
    return None


def main():
    print("=" * 60)
    print(" refinement: Arabic Script Refinement via Gemini")
    print("=" * 60)

    # 1. Read config variables FIRST
    model_name = get_config_value("REFINE_MODEL", "Pro")
    max_retries = int(get_config_value("FAILOVER_RETRY_LIMIT", "4"))
    switch_accounts = (
        get_config_value("SWITCH_ACCOUNTS_ENABLED", "false").strip().lower()
        in ("true", "1", "yes")
    )
    browser_type = get_config_value("BROWSER_TYPE", "chrome")
    profile_index = int(get_config_value("ACTIVE_PROFILE_INDEX", "1"))

    print(
        f"[CONFIG] Active Profile Index: {profile_index} | Browser: {browser_type} | Model: {model_name}"
    )

    # 2. Automatically verify or launch Chrome debug session using the configured profile index
    if not ensure_chrome_debug_session(browser_type, profile_index):
        print("Could not verify or start Chrome debugging session. Exiting.")
        return

    folder = get_latest_run_folder()
    if not folder:
        print("No youtube_runs folder found.")
        sys.exit(1)

    video_title = os.path.basename(os.path.normpath(folder))
    print(f"Processing: {video_title}")

    final_output = read_final_output(folder)
    paragraphs = split_paragraphs(final_output)
    print(f"Found {len(paragraphs)} paragraphs to refine.")

    refined_paragraphs = load_checkpoint(folder)
    start_index = len(refined_paragraphs)
    print(f"Resuming from paragraph {start_index + 1}.")

    if start_index >= len(paragraphs):
        print("All paragraphs already refined.")
        delete_checkpoint(folder)
        return

    with sync_playwright() as p:
        browser = p.chromium.connect_over_cdp(
            "http://127.0.0.1:9222", timeout=5000
        )
        context = browser.contexts[0]
        page = context.new_page()

        retries = 0
        current_index = start_index

        while current_index < len(paragraphs) and retries < max_retries:
            if current_index == start_index:
                if not setup_refinement_session(page, model_name):
                    print("[FATAL] Setup failed. Retrying with new chat...")
                    retries += 1
                    continue

            paragraph = paragraphs[current_index]
            print(
                f"\n[REFINE] Paragraph {current_index + 1}/{len(paragraphs)}..."
            )

            refined = refine_paragraph(
                page, paragraph, current_index + 1, len(paragraphs)
            )

            if refined and not is_safety_blocked(refined):
                refined_paragraphs.append(refined)
                save_checkpoint(folder, refined_paragraphs)
                print(
                    f"[OK] Paragraph {current_index + 1} refined ({len(refined)} chars)."
                )
                current_index += 1
                retries = 0
            else:
                retries += 1
                print(
                    f"[RETRY {retries}/{max_retries}] Paragraph {current_index + 1} failed. Starting fresh chat..."
                )

                if switch_accounts and retries >= max_retries:
                    print("[FAILOVER] Rotating account...")
                    profile_index = rotate_profile_index()
                    kill_cdp_chrome()
                    time.sleep(3)
                    launch_browser_with_profile(browser_type, profile_index)
                    time.sleep(5)
                    browser = p.chromium.connect_over_cdp(
                        "http://127.0.0.1:9222", timeout=5000
                    )
                    context = browser.contexts[0]
                    page = context.new_page()
                    retries = 0

                start_clean_gemini_chat(page)
                time.sleep(2)
                select_gemini_model(page, model_name)
                time.sleep(2)

        page.close()

    if current_index >= len(paragraphs):
        save_refined_script(folder, refined_paragraphs)
        delete_checkpoint(folder)
        print(f"\n{'=' * 60}")
        print(f" REFINEMENT COMPLETE: {video_title}")
        print(f"{'=' * 60}")
        send_telegram_notification(f"✅ Script refined: {video_title}")
    else:
        print(
            f"\n[PARTIAL] Refined {current_index}/{len(paragraphs)} paragraphs. Checkpoint saved."
        )
        send_telegram_notification(
            f"⚠️ Script refinement partial: {video_title} ({current_index}/{len(paragraphs)})"
        )

def verify_script_with_rubric(page, script_text):
    """Sends the refined script along with Payload 7.6 to verify quality."""
    with open("audit_rubric.md", "r", encoding="utf-8") as f:
        rubric_text = f.read()

    verification_prompt = (
        f"You are the Lead Script Doctor. Audit this script using the rubric below:\n\n"
        f"RUBRIC:\n{rubric_text}\n\n"
        f"SCRIPT TO AUDIT:\n{script_text}\n\n"
        f"Reply strictly with 'PASS' if it meets all 10 criteria, or list the specific failures."
    )
    # Send to Gemini to verify before proceeding to TTS
if __name__ == "__main__":
    main()