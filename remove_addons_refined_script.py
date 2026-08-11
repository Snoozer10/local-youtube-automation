import glob
import os
import re
import sys

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def get_target_file_path():
    """Locate refined_script.txt in the latest youtube_runs folder or current directory."""
    runs_path = "youtube_runs"
    if os.path.exists(runs_path):
        folders = glob.glob(os.path.join(runs_path, "*/"))
        if folders:
            latest_folder = max(folders, key=os.path.getmtime)
            target_txt = os.path.join(latest_folder, "refined_script.txt")
            if os.path.exists(target_txt):
                return target_txt, latest_folder

    # Fallback to current working directory
    if os.path.exists("refined_script.txt"):
        return "refined_script.txt", os.getcwd()

    return None, None


def clean_narrative_text(raw_text):
    """Deep cleaning function to strip all XML tags, thinking metadata, and ledger add-ons."""
    if not raw_text:
        return ""

    # 1. Remove XML blocks: <thinking>...</thinking> and <slang_ledger>...</slang_ledger>
    text = re.sub(r'<thinking>.*?</thinking>', '', raw_text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<thinking>.*', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<slang_ledger>.*?</slang_ledger>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<slang_ledger>.*', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<.*?>', '', text)  # Catch any stray XML tags

    # 2. Remove BLACKLIST / BRAINSTORM / DRAFT meta-blocks that leaked outside thinking tags
    text = re.sub(
        r'(?:BLACKLIST|BRAINSTORM|UPDATE LEDGER|DRAFT & VERIFY).*?(?:Tashkeel added[^\.\n]*[\.\s]*|verified[\.\s]*|words[\.\s]*|words\.)',
        '',
        text,
        flags=re.DOTALL | re.IGNORECASE
    )

    # 3. Process line-by-line to filter out residual metadata lines
    lines = text.splitlines()
    cleaned_lines = []
    
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue

        # Skip known metadata headers or English instruction lines
        if (
            line_str.startswith("BLACKLIST:")
            or line_str.startswith("BRAINSTORM:")
            or line_str.startswith("UPDATE LEDGER:")
            or line_str.startswith("DRAFT & VERIFY:")
            or line_str.startswith("Sentence ")
            or "slang terms per sentence" in line_str.lower()
            or "blacklisted words" in line_str.lower()
            or "tashkeel added" in line_str.lower()
            or line_str.startswith("🎬")
            or line_str.startswith("مستوى الحماس")
            or line_str.startswith("الهدف النفسي")
        ):
            continue

        # Remove inline English annotations like "(1 slang: فُتُوَّة)" or "(The Hook)"
        line_str = re.sub(r'\(\d+\s+slang:[^\)]*\)', '', line_str, flags=re.IGNORECASE)
        line_str = re.sub(r'\([A-Za-z0-9\s\-_,\.\'&]+\)', '', line_str)

        # Remove markdown bold/italics
        line_str = re.sub(r'\*\*(.*?)\*\*', r'\1', line_str)
        line_str = re.sub(r'\*(.*?)\*', r'\1', line_str)

        # Clean multiple spaces
        line_str = re.sub(r'\s+', ' ', line_str).strip()

        if line_str:
            cleaned_lines.append(line_str)

    # Reconstruct into paragraphs
    cleaned_text = "\n\n".join(cleaned_lines)
    return cleaned_text


def main():
    print("=" * 60)
    print("🧹 Script Add-ons Cleaner for Refined Scripts")
    print("=" * 60)

    txt_path, folder = get_target_file_path()

    if not txt_path or not os.path.exists(txt_path):
        print("❌ Error: Could not find 'refined_script.txt'.")
        sys.exit(1)

    print(f"📄 Found target file: {txt_path}")

    # Read raw content
    with open(txt_path, "r", encoding="utf-8") as f:
        raw_content = f.read()

    print(f"📊 Original file size: {len(raw_content)} characters.")

    # Clean script
    cleaned_content = clean_narrative_text(raw_content)

    if not cleaned_content:
        print("⚠️ Warning: Cleaned content is empty! Aborting overwrite to prevent data loss.")
        sys.exit(1)

    # Save cleaned TXT
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(cleaned_content)

    print(f"✅ Cleaned text saved to: {txt_path}")

    # Save cleaned DOCX if python-docx is available
    if HAS_DOCX and folder:
        docx_path = os.path.join(folder, "refined_script.docx")
        doc = Document()
        doc.add_heading("Refined Script (Cleaned)", level=1)
        for paragraph in cleaned_content.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        doc.save(docx_path)
        print(f"✅ Cleaned DOCX saved to: {docx_path}")

    print(f"📊 Cleaned file size: {len(cleaned_content)} characters.")
    print("🎉 Cleaning completed successfully!")


if __name__ == "__main__":
    main()